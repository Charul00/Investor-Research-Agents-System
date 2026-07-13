from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUTPUT_DIR = Path(__file__).resolve().parent / "generated" / "question-paper"
DOCX_PATH = OUTPUT_DIR / "atharva_infinity_classes_english_question_paper.docx"
PDF_PATH = OUTPUT_DIR / "atharva_infinity_classes_english_question_paper.pdf"
REPORT_WRITING_ROWS = 20
REPORT_WRITING_ROW_HEIGHT_IN = 0.31


def set_run_font(run, name="Arial", size=11, bold=None, color="000000"):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def style_font(style, name="Arial", size=11, bold=False, color="000000"):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        child = tc_mar.find(qn(f"w:{key}"))
        if child is None:
            child = OxmlElement(f"w:{key}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def configure_answer_cell(cell):
    set_cell_margins(cell)
    set_cell_borders(
        cell,
        left={"val": "nil"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(" ")


def clear_cell_borders(cell):
    set_cell_borders(
        cell,
        left={"val": "nil"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "nil"},
    )


def configure_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_header(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Atharva Sir's Infinity Classes")
    set_run_font(run, size=18, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    run = subtitle.add_run("English Grammar Question Paper")
    set_run_font(run, size=14, bold=True)

    marks = doc.add_paragraph()
    marks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marks.paragraph_format.space_after = Pt(12)
    run = marks.add_run("Total Marks: 25")
    set_run_font(run, size=11, bold=True)

    info = doc.add_table(rows=1, cols=4)
    configure_table(info, [Inches(1.1), Inches(2.2), Inches(0.7), Inches(2.5)])
    info.rows[0].height = Inches(0.32)
    info.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    labels = [("Name", 0), ("", 1), ("Date", 2), ("", 3)]
    for text, index in labels:
        cell = info.cell(0, index)
        clear_cell_borders(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        if text:
            run = paragraph.add_run(f"{text}:")
            set_run_font(run, bold=True)
        else:
            configure_answer_cell(cell)

    doc.add_paragraph()


def add_question_heading(doc, text):
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=16, bold=False)


def add_q1(doc):
    add_question_heading(doc, "Q1) Add a Question Tag (5 marks)")
    prompts = [
        "1. She wasn't tired",
        "2. We have finished the work",
        "3. You haven't met him",
        "4. He will call you",
        "5. You should study more",
    ]

    table = doc.add_table(rows=len(prompts), cols=2)
    configure_table(table, [Inches(4.85), Inches(1.65)])

    for row, prompt in zip(table.rows, prompts):
        row.height = Inches(0.35)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        question_cell, answer_cell = row.cells
        clear_cell_borders(question_cell)
        q_paragraph = question_cell.paragraphs[0]
        q_paragraph.paragraph_format.space_after = Pt(0)
        run = q_paragraph.add_run(prompt)
        set_run_font(run)
        configure_answer_cell(answer_cell)


def add_q2(doc):
    add_question_heading(doc, "Q2) Write the Structure and One Example of Each Tense (10 marks)")
    tenses = [
        "1. Simple Past Tense",
        "2. Future Perfect Tense",
        "3. Past Continuous Tense",
        "4. Present Perfect Continuous Tense",
        "5. Present Continuous Tense",
    ]

    for tense in tenses:
        item = doc.add_paragraph()
        item.paragraph_format.space_after = Pt(4)
        run = item.add_run(tense)
        set_run_font(run, bold=True)

        table = doc.add_table(rows=2, cols=2)
        configure_table(table, [Inches(1.15), Inches(5.35)])

        labels = ["Structure", "Example"]
        for row, label in zip(table.rows, labels):
            row.height = Inches(0.32)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            label_cell, answer_cell = row.cells
            clear_cell_borders(label_cell)
            label_paragraph = label_cell.paragraphs[0]
            label_paragraph.paragraph_format.space_after = Pt(0)
            label_run = label_paragraph.add_run(f"{label}:")
            set_run_font(label_run, bold=True)
            configure_answer_cell(answer_cell)

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)


def add_q3(doc):
    add_question_heading(doc, "Q3) Report Writing - Any One (5 marks)")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(4)
    run = intro.add_run("Write a report on any one of the following topics:")
    set_run_font(run)

    topics = [
        "1. New English School Celebrates 75th Independence Day",
        "2. A Terrible Road Accident",
    ]
    for topic in topics:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(topic)
        set_run_font(run)

    doc.add_paragraph()

    lines = doc.add_table(rows=REPORT_WRITING_ROWS, cols=1)
    configure_table(lines, [Inches(6.5)])
    for row in lines.rows:
        row.height = Inches(REPORT_WRITING_ROW_HEIGHT_IN)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cell = row.cells[0]
        configure_answer_cell(cell)


def add_q4(doc):
    doc.add_page_break()
    add_question_heading(doc, "Q4) Complete the Table: V1 - V2 - V3 (5 marks)")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run("Fill in the correct forms of V2 and V3.")
    set_run_font(run)

    table = doc.add_table(rows=11, cols=4)
    configure_table(table, [Inches(0.7), Inches(2.4), Inches(1.7), Inches(1.7)])

    header = ["No.", "V1", "V2", "V3"]
    for cell, text in zip(table.rows[0].cells, header):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, bold=True)
        set_cell_borders(
            cell,
            left={"val": "single", "sz": 8, "color": "DADCE0"},
            top={"val": "single", "sz": 8, "color": "DADCE0"},
            right={"val": "single", "sz": 8, "color": "DADCE0"},
            bottom={"val": "single", "sz": 8, "color": "DADCE0"},
        )

    verbs = [
        "Find",
        "Sink",
        "Wake",
        "Tear",
        "Hold",
        "Have",
        "Speak",
        "Ride",
        "Know",
        "Leave",
    ]

    border = {
        "left": {"val": "single", "sz": 8, "color": "DADCE0"},
        "top": {"val": "single", "sz": 8, "color": "DADCE0"},
        "right": {"val": "single", "sz": 8, "color": "DADCE0"},
        "bottom": {"val": "single", "sz": 8, "color": "DADCE0"},
    }

    for index, verb in enumerate(verbs, start=1):
        row = table.rows[index]
        row.height = Inches(0.34)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        values = [str(index), verb, "", ""]
        for cell, value, centered in zip(row.cells, values, [True, False, False, False]):
            set_cell_borders(cell, **border)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            if value:
                run = paragraph.add_run(value)
                set_run_font(run)


def build_question_paper():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    style_font(normal, size=11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    heading2 = doc.styles["Heading 2"]
    style_font(heading2, size=16, bold=False)
    heading2.paragraph_format.space_before = Pt(18)
    heading2.paragraph_format.space_after = Pt(6)

    add_header(doc)
    add_q1(doc)
    add_q2(doc)
    add_q3(doc)
    add_q4(doc)

    doc.save(DOCX_PATH)
    build_question_paper_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


def pdf_styles():
    return {
        "title": ParagraphStyle(
            "title",
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "center_note": ParagraphStyle(
            "center_note",
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "heading": ParagraphStyle(
            "heading",
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            spaceBefore=12,
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "body",
            fontName="Helvetica",
            fontSize=11,
            leading=14,
            spaceAfter=4,
        ),
        "item": ParagraphStyle(
            "item",
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            spaceAfter=4,
        ),
        "table_header": ParagraphStyle(
            "table_header",
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
        ),
    }


def q1_pdf_table(style):
    rows = [
        ["1. She wasn't tired", ""],
        ["2. We have finished the work", ""],
        ["3. You haven't met him", ""],
        ["4. He will call you", ""],
        ["5. You should study more", ""],
    ]
    table = Table(
        [[Paragraph(left, style), right] for left, right in rows],
        colWidths=[4.85 * inch, 1.65 * inch],
        rowHeights=[0.34 * inch] * len(rows),
    )
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("LINEBELOW", (1, 0), (1, -1), 1, colors.black),
            ]
        )
    )
    return table


def q2_pdf_tables(body_style, item_style):
    story = []
    tenses = [
        "1. Simple Past Tense",
        "2. Future Perfect Tense",
        "3. Past Continuous Tense",
        "4. Present Perfect Continuous Tense",
        "5. Present Continuous Tense",
    ]
    for tense in tenses:
        story.append(Paragraph(tense, item_style))
        table = Table(
            [
                [Paragraph("<b>Structure:</b>", body_style), ""],
                [Paragraph("<b>Example:</b>", body_style), ""],
            ],
            colWidths=[1.15 * inch, 5.35 * inch],
            rowHeights=[0.3 * inch, 0.3 * inch],
        )
        table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                    ("TOPPADDING", (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                    ("LINEBELOW", (1, 0), (1, -1), 1, colors.black),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 6))
    return story


def q3_pdf_table():
    table = Table(
        [[" "]] * REPORT_WRITING_ROWS,
        colWidths=[6.5 * inch],
        rowHeights=[REPORT_WRITING_ROW_HEIGHT_IN * inch] * REPORT_WRITING_ROWS,
    )
    table.setStyle(
        TableStyle(
            [
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("LINEBELOW", (0, 0), (0, -1), 1, colors.black),
            ]
        )
    )
    return table


def q4_pdf_table(body_style, header_style):
    rows = [["No.", "V1", "V2", "V3"]]
    verbs = ["Find", "Sink", "Wake", "Tear", "Hold", "Have", "Speak", "Ride", "Know", "Leave"]
    for index, verb in enumerate(verbs, start=1):
        rows.append([str(index), verb, "", ""])

    table = Table(
        rows,
        colWidths=[0.7 * inch, 2.4 * inch, 1.7 * inch, 1.7 * inch],
        rowHeights=[0.34 * inch] * len(rows),
        repeatRows=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.8, colors.HexColor("#555555")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (0, -1), "CENTER"),
                ("ALIGN", (1, 0), (-1, 0), "CENTER"),
            ]
        )
    )

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            if value:
                style = header_style if row_index == 0 else body_style
                table._cellvalues[row_index][col_index] = Paragraph(value, style)
    return table


def build_question_paper_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    story = [
        Paragraph("Atharva Sir's Infinity Classes", styles["title"]),
        Paragraph("English Grammar Question Paper", styles["subtitle"]),
        Paragraph("Total Marks: 25", styles["center_note"]),
    ]

    info = Table(
        [[Paragraph("<b>Name:</b>", styles["body"]), "", Paragraph("<b>Date:</b>", styles["body"]), ""]],
        colWidths=[1.1 * inch, 2.2 * inch, 0.7 * inch, 2.5 * inch],
        rowHeights=[0.32 * inch],
    )
    info.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("LINEBELOW", (1, 0), (1, 0), 1, colors.black),
                ("LINEBELOW", (3, 0), (3, 0), 1, colors.black),
            ]
        )
    )
    story.extend([info, Spacer(1, 12)])

    story.append(Paragraph("Q1) Add a Question Tag (5 marks)", styles["heading"]))
    story.append(q1_pdf_table(styles["body"]))

    story.append(Paragraph("Q2) Write the Structure and One Example of Each Tense (10 marks)", styles["heading"]))
    story.extend(q2_pdf_tables(styles["body"], styles["item"]))

    story.append(Paragraph("Q3) Report Writing - Any One (5 marks)", styles["heading"]))
    story.append(Paragraph("Write a report on any one of the following topics:", styles["body"]))
    story.append(Paragraph("1. New English School Celebrates 75th Independence Day", styles["body"]))
    story.append(Paragraph("2. A Terrible Road Accident", styles["body"]))
    story.append(Spacer(1, 6))
    story.append(q3_pdf_table())

    story.append(PageBreak())
    story.append(Paragraph("Q4) Complete the Table: V1 - V2 - V3 (5 marks)", styles["heading"]))
    story.append(Paragraph("Fill in the correct forms of V2 and V3.", styles["body"]))
    story.append(q4_pdf_table(styles["body"], styles["table_header"]))

    doc.build(story)


if __name__ == "__main__":
    build_question_paper()
